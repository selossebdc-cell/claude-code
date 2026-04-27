#!/usr/bin/env python3
"""Génère le CR d'audit Marion / L'Artisan Academy en .docx — Charte CS Consulting Stratégique"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# === COULEURS CHARTE CS BUSINESS ===
TERRACOTTA = RGBColor(0xD1, 0x7C, 0x61)
TERRE = RGBColor(0x43, 0x3F, 0x3C)
CUIVRE = RGBColor(0xAE, 0x75, 0x58)
SABLE = RGBColor(0xC1, 0x9E, 0x8A)
LIN = RGBColor(0xF3, 0xEA, 0xDA)
BLANC = RGBColor(0xFF, 0xFF, 0xFF)

LOGO_PATH = "/Users/cath/Library/CloudStorage/OneDrive-Personnel/01_Entreprise/05_Documents entreprise/LOGO CONSULTING STRATÉGIQUE/PNG/Catherine Selosse - Logo Seul_Logo seul - Terracotta.png"
OUTPUT_PATH = os.path.expanduser("~/Downloads/CR Audit Marion - Artisan Academy - CS Consulting Stratégique.docx")


def set_cell_shading(cell, color_hex):
    """Applique un fond de couleur à une cellule."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), val.get("sz", "4"))
            el.set(qn("w:color"), val.get("color", "D17C61"))
            el.set(qn("w:space"), "0")
            tcBorders.append(el)
    tcPr.append(tcBorders)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def add_styled_table(doc, headers, rows, header_bg="D17C61", header_fg=BLANC):
    """Crée un tableau stylé CS Consulting Stratégique."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = header_fg
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, header_bg)

    # Rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.color.rgb = TERRE
            run.font.name = "Calibri"
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F3EADA")

    # Spacing
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)

    return table


def add_heading_styled(doc, text, level=1):
    """Heading avec couleur Terracotta."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = TERRACOTTA
        run.font.name = "Calibri"
    h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    h.paragraph_format.space_after = Pt(8)
    return h


def add_subheading(doc, text):
    """Sous-titre en Terre gras."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = TERRE
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, bold=False, italic=False):
    """Paragraphe corps de texte."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = TERRE
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, bold_prefix=None):
    """Puce avec optionnel préfixe en gras."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_b.font.color.rgb = TERRE
        run_b.font.name = "Calibri"
        run_n = p.add_run(text)
        run_n.font.size = Pt(11)
        run_n.font.color.rgb = TERRE
        run_n.font.name = "Calibri"
    else:
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = TERRE
        run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(3)
    return p


def add_separator(doc):
    """Ligne de séparation terracotta."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Add a bottom border to the paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D17C61")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_highlight_box(doc, text):
    """Encadré avec fond Lin pour mise en avant."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = TERRACOTTA
    run.font.name = "Calibri"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, "F3EADA")
    border_style = {"val": "single", "sz": "8", "color": "D17C61"}
    set_cell_borders(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    # table width
    for cell in table.rows[0].cells:
        cell.width = Inches(6)


# === CONSTRUCTION DU DOCUMENT ===
doc = Document()

# --- Marges ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Style par défaut ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = TERRE

# ==========================================
# PAGE DE COUVERTURE
# ==========================================

# Logo
if os.path.exists(LOGO_PATH):
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run()
    run_logo.add_picture(LOGO_PATH, width=Inches(1.5))

# Espace
doc.add_paragraph()
doc.add_paragraph()

# Titre principal
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_t = p_title.add_run("COMPTE-RENDU D'AUDIT")
run_t.font.size = Pt(28)
run_t.font.color.rgb = TERRACOTTA
run_t.font.name = "Calibri"
run_t.bold = True

p_subtitle = doc.add_paragraph()
p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_s = p_subtitle.add_run("& PRÉCONISATIONS")
run_s.font.size = Pt(22)
run_s.font.color.rgb = SABLE
run_s.font.name = "Calibri"

doc.add_paragraph()

# Séparateur
add_separator(doc)

doc.add_paragraph()

# Infos client
for line in [
    ("Marion — L'Artisan Academy", True, Pt(16), TERRE),
    ("", False, Pt(8), TERRE),
    ("Date de l'audit : 24 février 2026", False, Pt(12), CUIVRE),
    ("", False, Pt(6), TERRE),
    ("Préparé par Catherine Selosse", False, Pt(12), CUIVRE),
    ("CS Consulting Stratégique", False, Pt(11), SABLE),
    ("catherine@csbusiness.fr", False, Pt(11), SABLE),
]:
    text, bold, size, color = line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = size
    run.font.color.rgb = color
    run.font.name = "Calibri"
    run.bold = bold
    p.paragraph_format.space_after = Pt(2)

# Saut de page
doc.add_page_break()

# ==========================================
# OUVERTURE
# ==========================================

add_heading_styled(doc, "Merci Marion,", level=1)

add_body(doc,
    "Merci pour cet échange, ta clarté et ta transparence sur tes besoins. "
    "Tu as bâti quelque chose d'impressionnant : plus de 2 000 élèves, "
    "500 en accompagnement actif, une équipe de 20 personnes, 9 marraines "
    "qui font vivre ta communauté, et une pédagogie structurée en processus "
    "— pas une simple boîte à outils."
)

add_body(doc,
    "Ce que tu cherches aujourd'hui, c'est de passer un cap : mettre l'IA "
    "au service de tes élèves pour qu'elles avancent plus vite, plus sereinement, "
    "et à plus grande échelle. Ton ambition est claire, tes besoins sont précis."
)

add_body(doc,
    "Voici ce que j'ai compris et ce que je te préconise.",
    bold=True
)

add_separator(doc)

# ==========================================
# 1. SYNTHÈSE DES BESOINS
# ==========================================

add_heading_styled(doc, "1. Synthèse de tes besoins", level=1)

add_body(doc,
    "Suite à notre échange, j'ai identifié 3 pôles de besoins, par ordre de priorité :"
)

# --- Besoin prioritaire ---
add_subheading(doc, "Besoin prioritaire — Un agent IA pour tes élèves (lancement juin)")

add_body(doc,
    "Tu souhaites un assistant IA intégré à ton école qui puisse accompagner "
    "tes élèves sur plusieurs dimensions :"
)

add_styled_table(doc,
    ["Fonctionnalité", "Ce que tu veux concrètement"],
    [
        ["Navigation dans la formation",
         "L'IA guide chaque élève vers le bon module selon son niveau et sa situation (débutante vs déjà lancée)"],
        ["Posture de coach",
         "L'IA pose les bonnes questions, audite la situation de la créatrice, et l'oriente — pas juste un moteur de recherche"],
        ["Génération de livrables",
         "Templates pré-remplis par l'IA (calendrier éditorial, rétroplanning, emails...) à partir de questions guidées"],
        ["Correction de cahiers d'exercices",
         "L'IA analyse les réponses des PDF remplissables et donne un retour personnalisé selon tes critères pédagogiques"],
        ["Matching de binômes",
         "Création de binômes de travail intelligents basés sur les profils (pas aléatoire)"],
        ["Analyse de performances",
         "L'IA analyse les KPI boutique (visibilité, trafic, commandes, panier moyen) et identifie les pistes d'amélioration"],
    ]
)

doc.add_paragraph()

add_body(doc, "Tes contraintes :", bold=True)
add_bullet(doc, "Intégration simple : Discord (accompagnement) et/ou SchoolMaker (formation) — pas d'outil externe complexe")
add_bullet(doc, "Différenciation entre la formule Accompagnement (500 élèves) et Autonomie (1 500+)")
add_bullet(doc, "Maîtrise des coûts de tokens sur un volume de 2 000+ élèves")
add_bullet(doc, "Propriété des données et de la solution")

# --- Besoin secondaire ---
add_subheading(doc, "Besoin secondaire — Base de connaissances interne (pas urgent)")

add_body(doc,
    "Capitaliser 4 ans de propriété intellectuelle pour que ton équipe de "
    "20 personnes accède facilement aux infos sans toujours te solliciter. "
    "Un agent interne par pôle (marketing, admin, etc.)."
)

# --- Idée business ---
add_subheading(doc, "Idée business — Livrables payants (à explorer)")

add_body(doc,
    "Micro-pricing sur les livrables générés par l'IA (3 €, 5 €, 10 €) pour créer "
    "un nouveau levier de revenus récurrent au sein de l'école."
)

add_body(doc,
    "Comme tu l'as dit : « Si je mets du budget pour développer une solution custom, "
    "il faut que ce soit une solution qui rapporte de l'argent. »",
    italic=True
)

add_separator(doc)

# ==========================================
# 2. RECOMMANDATION : ARIA
# ==========================================

add_heading_styled(doc, "2. Ma recommandation : ARIA comme brique mère", level=1)

add_body(doc,
    "Après avoir écouté l'ensemble de tes besoins et ton contexte, "
    "ma recommandation est claire et honnête :"
)

add_highlight_box(doc, "Prends ARIA comme solution principale.")

doc.add_paragraph()

add_body(doc, "ARIA couvre la majorité de tes besoins fondamentaux :")
add_bullet(doc, "Navigation dans la formation")
add_bullet(doc, "Posture coach / guidage des élèves")
add_bullet(doc, "Intégration avec tes plateformes")
add_bullet(doc, "Gestion de la volumétrie et des tokens incluse dans leur forfait")
add_bullet(doc, "Mise en place rapide, adaptée à ton lancement de juin")

doc.add_paragraph()

add_body(doc,
    "C'est une solution éprouvée, recommandée dans ton écosystème, et conçue "
    "spécifiquement pour les infopreneurs. Je ne suis pas partisane de réinventer "
    "la roue quand quelque chose fonctionne bien et que l'équipe derrière est solide. "
    "ARIA est cette base solide."
)

add_body(doc,
    "Mon rôle : te proposer les briques qui n'existent pas chez ARIA.",
    bold=True
)

add_separator(doc)

# ==========================================
# 3. BRIQUES COMPLÉMENTAIRES
# ==========================================

add_heading_styled(doc, "3. Briques complémentaires — Ce que je propose", level=1)

add_body(doc,
    "Voici les fonctionnalités sur mesure que je peux développer en complément d'ARIA. "
    "Ce sont les éléments qui nécessitent un développement custom parce qu'ils n'existent "
    "pas en standard."
)

# --- Brique 1 ---
add_heading_styled(doc, "Brique 1 — Plateforme de livrables payants", level=2)

add_body(doc,
    "Objectif : Transformer l'IA en générateur de revenus, pas juste un outil de support.",
    bold=True
)

add_body(doc, "Ce que ça inclut :")
add_bullet(doc, " intégré (Stripe) pour chaque livrable généré par l'IA", bold_prefix="Système de micro-paiement")
add_bullet(doc, " configurables : calendrier éditorial, rétroplanning lancement, séquences emails, fiches produit, etc.", bold_prefix="Catalogue de livrables")
add_bullet(doc, " pour chaque livrable : l'élève répond à 5-10 questions, l'IA génère le document personnalisé", bold_prefix="Questionnaire guidé")
add_bullet(doc, " des ventes : quels livrables se vendent, combien, à qui", bold_prefix="Dashboard de suivi")
add_bullet(doc, " pensée dès le départ pour le paiement, même si la monétisation se fait dans un second temps", bold_prefix="Architecture")

doc.add_paragraph()
add_body(doc, "Pourquoi c'est stratégique :", bold=True)
add_bullet(doc, "Crée un revenu récurrent indépendant des lancements")
add_bullet(doc, "Valorise ta méthodologie sous un nouveau format")
add_bullet(doc, "L'IA produit, tu encaisses — scalable par nature")

# --- Brique 2 ---
add_heading_styled(doc, "Brique 2 — Correction intelligente des cahiers d'exercices", level=2)

add_body(doc,
    "Objectif : Débloquer les créatrices qui n'osent pas avancer par manque de confiance.",
    bold=True
)

add_body(doc, "Ce que ça inclut :")
add_bullet(doc, " et d'analyse des PDF remplissables (Acrobat Reader)", bold_prefix="Système de lecture")
add_bullet(doc, " personnalisées par cahier d'exercice, selon tes critères pédagogiques", bold_prefix="Grilles de correction")
add_bullet(doc, " à l'élève : ce qui est bien, ce qui peut être amélioré, la prochaine étape", bold_prefix="Retour structuré")
add_bullet(doc, " bienveillant et encourageant — programmé selon ta posture de formatrice", bold_prefix="Ton")

doc.add_paragraph()
add_body(doc, "Pourquoi c'est stratégique :", bold=True)
add_bullet(doc, "C'est un travail monumental, impossible à faire humainement sur 2 000 élèves")
add_bullet(doc, "C'est le déblocage n°1 : tes élèves ont besoin d'être validées pour continuer")
add_bullet(doc, "Ça réduit la charge de tes marraines sur les questions récurrentes")

# --- Brique 3 ---
add_heading_styled(doc, "Brique 3 — Matching intelligent de binômes", level=2)

add_body(doc,
    "Objectif : Reproduire à grande échelle ce qui se passe naturellement quand les créatrices se rencontrent en live.",
    bold=True
)

add_body(doc, "Ce que ça inclut :")
add_bullet(doc, " de profiling : secteur, niveau, objectifs, personnalité, disponibilités", bold_prefix="Questionnaire")
add_bullet(doc, " de matching basé sur la complémentarité (pas la similarité)", bold_prefix="Algorithme")
add_bullet(doc, " de binômes avec justification (« Vous êtes complémentaires parce que... »)", bold_prefix="Suggestions")
add_bullet(doc, " possible pour renouveler les dynamiques", bold_prefix="Rotation")

doc.add_paragraph()
add_body(doc, "Pourquoi c'est stratégique :", bold=True)
add_bullet(doc, "Tu l'as observé : quand les bonnes personnes se trouvent, c'est très puissant")
add_bullet(doc, "Ça renforce l'engagement et la rétention dans la formule Accompagnement")
add_bullet(doc, "C'est un vrai différenciateur par rapport aux autres écoles en ligne")

# --- Brique 4 ---
add_heading_styled(doc, "Brique 4 (optionnelle) — Analyse de performances boutique", level=2)

add_body(doc,
    "Objectif : Donner à chaque créatrice un diagnostic chiffré de sa boutique.",
    bold=True
)

add_body(doc, "Ce que ça inclut :")
add_bullet(doc, " structuré ou connexion aux données boutique (selon la plateforme e-commerce)", bold_prefix="Questionnaire")
add_bullet(doc, " des KPI clés : visibilité, engagement, trafic, taux de conversion, panier moyen", bold_prefix="Analyse")
add_bullet(doc, " automatiques reliées à ta formation (« Tu as un problème de trafic → va voir la leçon X »)", bold_prefix="Recommandations")
add_bullet(doc, " dans le temps pour mesurer la progression", bold_prefix="Suivi")

add_separator(doc)

# ==========================================
# 4. COMMENT ON PROCÈDE
# ==========================================

add_heading_styled(doc, "4. Comment on procède", level=1)

# Étape 1
add_subheading(doc, "Étape 1 — Tu avances avec ARIA (dès maintenant)")
add_body(doc,
    "N'attends pas ma proposition pour lancer ARIA. C'est ta brique mère, "
    "elle couvre tes besoins fondamentaux, et ton lancement de juin n'attend pas. "
    "Lance la mise en place."
)

# Étape 2
add_subheading(doc, "Étape 2 — Je fais le point avec mon développeur technique")
add_body(doc,
    "Je vais analyser en détail la faisabilité technique de chaque brique complémentaire. "
    "Mon développeur, spécialisé en cybersécurité, va chiffrer précisément le développement, "
    "les intégrations nécessaires, et les aspects sécurité (protection des données élèves, "
    "hébergement, etc.)."
)

# Étape 3
add_subheading(doc, "Étape 3 — Je te reviens avec le chiffrage")
add_body(doc, "Dès que le chiffrage est finalisé, je t'envoie la proposition complète avec :")
add_bullet(doc, "Le prix de chaque brique (format à tiroirs — tu choisis ce que tu veux)")
add_bullet(doc, "Les délais de développement")
add_bullet(doc, "Les options d'hébergement et de maintenance")
add_bullet(doc, "Les coûts récurrents éventuels")

# Étape 4
add_subheading(doc, "Étape 4 — On en discute (avec ton mari si besoin)")
add_body(doc,
    "On se refait un point pour valider ensemble les priorités, le planning, et les choix "
    "techniques. Ton mari, en tant que DAF et responsable IA, sera le bienvenu dans cet échange."
)

add_separator(doc)

# ==========================================
# 5. MES ENGAGEMENTS
# ==========================================

add_heading_styled(doc, "5. Ce que je m'engage à faire", level=1)

add_bullet(doc, " Si ARIA couvre un besoin, je te le dirai. Je ne développe pas ce qui existe déjà.", bold_prefix="Être transparente.")
add_bullet(doc, " Code, données, documentation — tout t'appartient.", bold_prefix="Te livrer une solution dont tu es propriétaire.")
add_bullet(doc, " Hébergement dédié, couche cybersécurité validée par mon développeur.", bold_prefix="Sécuriser tes données élèves.")
add_bullet(doc, " Tuto, formation, documentation — ton mari ou toi pouvez reprendre la main.", bold_prefix="Te rendre autonome.")

add_separator(doc)

# ==========================================
# CLOSING
# ==========================================

doc.add_paragraph()

add_body(doc,
    "Je reviens vers toi dès que j'ai le chiffrage de mon développeur.",
    bold=True
)
add_body(doc,
    "En attendant, fonce avec ARIA pour ton lancement de juin."
)
add_body(doc,
    "Et on se voit à Bordeaux début avril !"
)

doc.add_paragraph()
doc.add_paragraph()

# Signature
p_sig1 = doc.add_paragraph()
p_sig1.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p_sig1.add_run("Catherine Selosse")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = TERRACOTTA
run.font.name = "Calibri"

p_sig2 = doc.add_paragraph()
run = p_sig2.add_run("CS Consulting Stratégique")
run.font.size = Pt(11)
run.font.color.rgb = CUIVRE
run.font.name = "Calibri"

p_sig3 = doc.add_paragraph()
run = p_sig3.add_run("catherine@csbusiness.fr")
run.font.size = Pt(11)
run.font.color.rgb = SABLE
run.font.name = "Calibri"

# ==========================================
# SAUVEGARDE
# ==========================================

doc.save(OUTPUT_PATH)
print(f"✅ Document généré : {OUTPUT_PATH}")
