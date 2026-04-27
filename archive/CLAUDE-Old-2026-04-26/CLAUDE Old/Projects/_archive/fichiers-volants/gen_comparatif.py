#!/usr/bin/env python3
"""Génère le comparatif plateformes Aurélia en Word avec la charte CS Business."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# === COULEURS CS BUSINESS ===
TERRACOTTA = RGBColor(0xD1, 0x7C, 0x61)
TERRE = RGBColor(0x43, 0x3F, 0x3C)
CUIVRE = RGBColor(0xAE, 0x75, 0x58)
SABLE = RGBColor(0xC1, 0x9E, 0x8A)
LIN = RGBColor(0xF3, 0xEA, 0xDA)
BLANC = RGBColor(0xFF, 0xFF, 0xFF)

# === POLICES (substitution Google: Playfair Display titres, Cormorant Garamond corps) ===
FONT_TITRE = "Georgia"  # fallback universel serif élégant
FONT_CORPS = "Georgia"

doc = Document()

# --- Styles de base ---
style = doc.styles["Normal"]
style.font.name = FONT_CORPS
style.font.size = Pt(10)
style.font.color.rgb = TERRE
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

# Marges
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_cell_shading(cell, color_hex):
    """Applique une couleur de fond à une cellule."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_heading(text, level=1):
    """Ajoute un titre stylé CS Business."""
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.font.size = Pt(18)
        run.font.color.rgb = TERRACOTTA
        run.font.name = FONT_TITRE
        run.bold = True
        # Ligne séparatrice
        sep = doc.add_paragraph()
        sep.paragraph_format.space_after = Pt(6)
        sep_run = sep.add_run("─" * 60)
        sep_run.font.color.rgb = SABLE
        sep_run.font.size = Pt(8)
    elif level == 2:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.color.rgb = TERRE
        run.font.name = FONT_TITRE
        run.bold = True
    elif level == 3:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = CUIVRE
        run.font.name = FONT_TITRE
        run.bold = True
    return p


def add_body(text, bold=False, italic=False, color=None):
    """Ajoute un paragraphe corps."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_CORPS
    run.font.size = Pt(10)
    run.font.color.rgb = color or TERRE
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(text, bold_prefix="", indent_level=0):
    """Ajoute un point bullet."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3 + indent_level * 0.3)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.name = FONT_CORPS
        run_b.font.size = Pt(10)
        run_b.font.color.rgb = TERRE
        run_b.bold = True
        run_t = p.add_run(text)
        run_t.font.name = FONT_CORPS
        run_t.font.size = Pt(10)
        run_t.font.color.rgb = TERRE
    else:
        run = p.add_run("•  " + text)
        run.font.name = FONT_CORPS
        run.font.size = Pt(10)
        run.font.color.rgb = TERRE
    return p


def add_table(headers, rows, highlight_col=None):
    """Ajoute un tableau stylé CS Business."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-têtes
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = FONT_CORPS
        run.font.size = Pt(9)
        run.font.color.rgb = BLANC
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "D17C61")

    # Lignes
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = FONT_CORPS
            run.font.size = Pt(9)
            run.font.color.rgb = TERRE
            if c == 0:
                run.bold = True
            if highlight_col is not None and c == highlight_col:
                run.font.color.rgb = TERRACOTTA
                run.bold = True
            # Alternance fond
            bg = "F3EADA" if r % 2 == 0 else "FFFFFF"
            set_cell_shading(cell, bg)

    # Bordures fines
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="C19E8A"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="C19E8A"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="C19E8A"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="C19E8A"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="C19E8A"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="C19E8A"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_paragraph()  # espace après tableau
    return table


def add_quote(text, author=""):
    """Ajoute une citation stylée."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f'"{text}"')
    run.font.name = FONT_CORPS
    run.font.size = Pt(10)
    run.font.color.rgb = CUIVRE
    run.italic = True
    if author:
        run2 = p.add_run(f"\n— {author}")
        run2.font.name = FONT_CORPS
        run2.font.size = Pt(9)
        run2.font.color.rgb = SABLE


def add_winner(text):
    """Ajoute une ligne 'gagnant'."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(f"Gagnant : {text}")
    run.font.name = FONT_TITRE
    run.font.size = Pt(10)
    run.font.color.rgb = TERRACOTTA
    run.bold = True


# =============================================================================
# PAGE DE COUVERTURE
# =============================================================================

# Espacement haut
for _ in range(6):
    doc.add_paragraph()

# Titre principal
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Comparatif Technique\ndes Plateformes de Formation")
run.font.name = FONT_TITRE
run.font.size = Pt(28)
run.font.color.rgb = TERRACOTTA
run.bold = True

# Ligne déco
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─────────────────")
run.font.color.rgb = SABLE
run.font.size = Pt(14)

# Sous-titre
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
run = p.add_run("Kajabi  •  Circle.so  •  Skool")
run.font.name = FONT_CORPS
run.font.size = Pt(14)
run.font.color.rgb = TERRE

# Info client
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prepare pour Aurelia DELSOL")
run.font.name = FONT_CORPS
run.font.size = Pt(12)
run.font.color.rgb = CUIVRE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Fevrier 2026")
run.font.name = FONT_CORPS
run.font.size = Pt(11)
run.font.color.rgb = SABLE

# Espacement bas
for _ in range(4):
    doc.add_paragraph()

# Signature
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Catherine Selosse — Consulting Strategique")
run.font.name = FONT_TITRE
run.font.size = Pt(11)
run.font.color.rgb = TERRE
run.bold = True

# Saut de page
doc.add_page_break()

# =============================================================================
# RESUME EXECUTIF
# =============================================================================

add_styled_heading("Resume Executif", 1)

add_styled_heading("Contexte client", 2)
add_bullet("CA actuel : 150 000 EUR (+40% croissance)")
add_bullet("Uscreen actuel : 6 000 EUR/an + 1 000 EUR stockage + 1 EUR/utilisateur (non scalable)")
add_bullet("500+ videos a migrer")
add_bullet("Besoins : Cours, communaute, lives, calendrier, vente de produits")
add_bullet("Objectif multilingue : Francais/Anglais a terme")
add_bullet("Equipe : Aurelia + Laurie (assistante) + 3 intervenants a integrer")
add_bullet("Nouveau modele : 299 EUR pour 3 mois (vs 17 EUR/mois avant)")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Recommandation principale : Circle.so (Plan Business)")
run.font.name = FONT_TITRE
run.font.size = Pt(13)
run.font.color.rgb = TERRACOTTA
run.bold = True

# =============================================================================
# TABLEAU COMPARATIF SYNTHETIQUE
# =============================================================================

add_styled_heading("Tableau Comparatif Synthetique", 1)

add_table(
    ["Critere", "Kajabi", "Circle", "Skool", "Note"],
    [
        ["Prix/mois", "179-249 EUR", "89-199 EUR", "99 EUR", "Circle meilleur rapport qualite/prix"],
        ["UX/Design", "★★★", "★★★★★", "★★★★★", "Kajabi juge fragmente"],
        ["UX Mobile", "★★★", "★★★★", "★★★★★", "Skool tres intuitif mobile"],
        ["Communaute", "★★★", "★★★★★", "★★★★★", "Circle et Skool excellents"],
        ["Cours avances", "★★★★★", "★★★★", "★★★", "Kajabi le plus complet"],
        ["Marketing/Funnels", "★★★★★", "★★★", "★", "Kajabi seul tout-en-un"],
        ["Scalabilite", "★★★★", "★★★★★", "★★★★", "Pas de frais par utilisateur"],
        ["Multilingue", "★★★★", "★★★", "★★", "Kajabi meilleur support natif"],
        ["Migration videos", "Manuel", "Service gratuit*", "Manuel", "*Plan Business annuel"],
        ["Videos externes", "Oui", "Oui", "Oui", "Tous supportent l'embed"],
    ],
    highlight_col=2,  # Circle en surbrillance
)

# =============================================================================
# 1. KAJABI
# =============================================================================

doc.add_page_break()
add_styled_heading("1. Kajabi — L'Ecosysteme Tout-en-Un", 1)

add_styled_heading("Tarification (Janvier 2026)", 2)
add_table(
    ["Plan", "Prix/mois (annuel)", "Produits", "Contacts", "Membres actifs"],
    [
        ["Kickstarter", "89 EUR (71 EUR)", "1", "250", "50"],
        ["Basic", "179 EUR (143 EUR)", "5", "2 500", "1 000"],
        ["Growth", "249 EUR (199 EUR)", "50", "25 000", "10 000"],
        ["Pro", "399 EUR (319 EUR)", "100", "100 000", "20 000"],
    ],
)

add_body("Note : Kajabi a augmente ses prix en janvier 2026 pour la premiere fois en 10 ans.", italic=True, color=CUIVRE)

add_styled_heading("Points Forts", 2)
add_styled_heading("Fonctionnalites completes", 3)
add_bullet("Creation de cours avancee (quiz, evaluations, certificats, graphiques de retention)")
add_bullet("Website builder integre avec templates professionnels")
add_bullet("Email marketing natif avec sequences et automations")
add_bullet("Pipelines/Funnels de vente pre-configures")
add_bullet("Coaching individuel et groupe (cohorts)")
add_bullet("Hebergement video illimite inclus")
add_bullet("Nouvelle fonctionnalite : Traduction et transcription video automatique")

add_styled_heading("Pour Aurelia specifiquement", 3)
add_bullet("Support multilingue plus mature")
add_bullet("Tout-en-un = moins d'outils a gerer")
add_bullet("Professionnalisme percu pour ses offres premium (299 EUR)")

add_styled_heading("Points Faibles", 2)
add_styled_heading("UX et Interface", 3)
add_bullet('Interface jugee "cluttered" (encombree) par de nombreux utilisateurs')
add_bullet("Apps separees pour cours et communaute = experience fragmentee")
add_bullet('Communaute 2.0 consideree "under-baked" par les experts')
add_bullet("Courbe d'apprentissage significative")

add_styled_heading("Couts caches", 3)
add_bullet("Si utilisation de Stripe au lieu de Kajabi Payments : +2% de frais")
add_bullet("Plans tres limites en entree de gamme")
add_bullet("Basic = seulement 5 produits et 1 000 membres actifs")

add_quote(
    "La fonctionnalite communaute de Kajabi est 'under-baked' - aucun des utilisateurs haut de gamme que j'ai consultes ne l'utilise reellement.",
    "Jordan Godbey, Community Strategist",
)

add_styled_heading("Cout Estime pour Aurelia", 2)
add_body("Plan recommande : Growth a 199 EUR/mois (annuel) = 2 388 EUR/an", bold=True)
add_bullet("50 produits (suffisant pour ses programmes)")
add_bullet("10 000 membres actifs")
add_bullet("Automations avancees")
add_bullet("Programme d'affiliation inclus")
add_body("Economie vs Uscreen : ~3 600 EUR/an", bold=True, color=TERRACOTTA)

# =============================================================================
# 2. CIRCLE.SO
# =============================================================================

doc.add_page_break()
add_styled_heading("2. Circle.so — Le Specialiste Communaute+Cours", 1)

add_styled_heading("Tarification (2026)", 2)
add_table(
    ["Plan", "Prix/mois (annuel)", "Membres", "Stockage", "Frais transaction"],
    [
        ["Professional", "89 EUR", "Illimites", "200 GB", "2%"],
        ["Business", "199 EUR", "Illimites", "500 GB", "1%"],
        ["Enterprise", "399 EUR", "Illimites", "1 TB", "0.5%"],
        ["Plus", "~30 000 EUR/an", "Illimites", "Illimite", "Sur mesure"],
    ],
    highlight_col=0,
)

add_styled_heading("Points Forts", 2)

add_styled_heading("UX Exceptionnelle", 3)
add_bullet("Interface moderne inspiree de Slack/Notion")
add_bullet("Navigation fluide et intuitive")
add_bullet("Mobile-friendly avec app native iOS/Android")
add_bullet("Personnalisation avancee du branding (CSS custom des le plan Professional)")

add_styled_heading("Communaute Premium", 3)
add_bullet('"Spaces" flexibles et organisables')
add_bullet("Profils membres riches favorisant le networking")
add_bullet("Gamification avec niveaux et recompenses")
add_bullet("Lives integres sans outil tiers")
add_bullet("Chat en temps reel et messagerie directe")

add_styled_heading("Cours integres", 3)
add_bullet("Course builder inclus (plan Professional+)")
add_bullet("Progression des etudiants trackee")
add_bullet("Drip content disponible")

add_styled_heading("Services de Migration GRATUITS", 3)
add_bullet("Migration paiements : Transfert des abonnements actifs depuis Uscreen", bold_prefix="")
add_bullet("Migration cours : Gratuite pour plans Business annuels et superieurs", bold_prefix="")
add_bullet("Support dedie pour la migration")

add_styled_heading("Pour Aurelia specifiquement", 3)
add_bullet("Migration assistee = enorme gain de temps pour ses 500+ videos", bold_prefix="")
add_bullet("Communaute premium alignee avec son positionnement (299 EUR)")
add_bullet("Possibilite d'integrer ses 3 intervenants comme admins/moderateurs")
add_bullet("Workflow automations (plan Business) pour deleguer a Laurie")

add_styled_heading("Points Faibles", 2)
add_bullet("Pas de website builder (besoin d'un site externe ou landing pages)")
add_bullet("Pas de funnels de vente integres")
add_bullet("Email marketing basique (broadcasts et drips simples uniquement)")
add_bullet("Pas de quiz/certificats avances comme Kajabi")
add_bullet("Frais de transaction sur tous les plans (0.5% a 2%)")
add_bullet("Necessitera probablement un outil email externe (Mailerlite, ConvertKit...)")

add_styled_heading("Cout Estime pour Aurelia", 2)
add_body("Plan recommande : Business a 199 EUR/mois (annuel) = 2 388 EUR/an", bold=True)
add_bullet("Workflows/Automations inclus")
add_bullet("Migration cours gratuite")
add_bullet("1% transaction fees seulement")
add_bullet("Stockage 500 GB suffisant")

add_styled_heading("Couts additionnels potentiels", 3)
add_bullet("Outil email : ~30-50 EUR/mois")
add_bullet("Landing page builder : ~30 EUR/mois (ou utiliser son site existant)")
add_body("Cout total estime : ~3 300 EUR/an", bold=True)
add_body("Economie vs Uscreen : ~3 700 EUR/an + pas de frais par utilisateur", bold=True, color=TERRACOTTA)

# =============================================================================
# 3. SKOOL
# =============================================================================

doc.add_page_break()
add_styled_heading("3. Skool — La Simplicite Gamifiee", 1)

add_styled_heading("Tarification (2026)", 2)
add_table(
    ["Plan", "Prix/mois", "Admins", "Frais transaction"],
    [
        ["Hobby", "9 EUR", "1", "10%"],
        ["Pro", "99 EUR", "Illimites", "2.9%"],
    ],
)

add_styled_heading("Points Forts", 2)

add_styled_heading("UX Inegalee", 3)
add_bullet("Interface la plus simple et intuitive du marche")
add_bullet("Ressemble a Facebook = familiarite immediate")
add_bullet("Meilleurs taux d'engagement compares a Kajabi et Circle")
add_bullet("App mobile excellente")

add_styled_heading("Gamification Native", 3)
add_bullet("Points, niveaux, badges, leaderboards")
add_bullet("Deverrouillage de contenu par niveau")
add_bullet("Encourage l'engagement et la retention")

add_styled_heading("Communaute decouvrable", 3)
add_bullet("Moteur de recherche de communautes (comme Facebook Groups)")
add_bullet("Acquisition organique potentielle de nouveaux membres")

add_styled_heading("Pour Aurelia", 3)
add_bullet("Ideal si la communaute est au coeur de son modele")
add_bullet("Prix tres competitif (99 EUR pour tout)")
add_bullet("Membres WhatsApp pourraient facilement migrer")

add_styled_heading("Points Faibles", 2)
add_bullet("PAS de quiz, evaluations, ou certificats", bold_prefix="")
add_bullet("PAS de personnalisation du branding (tous les Skool se ressemblent)", bold_prefix="")
add_bullet("PAS d'email marketing (necessite outil externe)", bold_prefix="")
add_bullet("PAS de website builder ou funnels", bold_prefix="")
add_bullet("Maximum 10 categories de discussion")
add_bullet("Pas de sous-communautes (1 groupe = 1 abonnement de 99 EUR)")

add_styled_heading("Pour Aurelia specifiquement", 3)
add_bullet("Peu adapte a son positionnement premium (299 EUR)")
add_bullet("Migration des 500 videos = 100% manuelle")
add_bullet("Integration de 3 intervenants compliquee (pas de roles granulaires)")
add_bullet("Multilingue non supporte nativement")

add_styled_heading("Cout Estime pour Aurelia", 2)
add_body("Plan recommande : Pro a 99 EUR/mois = 1 188 EUR/an", bold=True)
add_styled_heading("Couts additionnels necessaires", 3)
add_bullet("Outil email : ~50 EUR/mois = 600 EUR/an")
add_bullet("Landing pages/Funnels : ~50 EUR/mois = 600 EUR/an")
add_bullet("Calendly Pro : ~15 EUR/mois = 180 EUR/an")
add_body("Cout total estime : ~2 500-3 000 EUR/an", bold=True)
add_body("Attention : Effort de migration manuel considerable + limitations fonctionnelles", italic=True, color=CUIVRE)

# =============================================================================
# 4. HEBERGEMENT VIDEO EXTERNE
# =============================================================================

doc.add_page_break()
add_styled_heading("4. Hebergement Video Externe — La Strategie Economique", 1)

add_styled_heading("Pourquoi heberger les videos separement ?", 2)
add_bullet("Independence : Pas enfermee dans une plateforme")
add_bullet("Economies : Potentiellement bien moins cher")
add_bullet("Migration facilitee : Les videos restent si la plateforme change")
add_bullet("Performance : CDN dedie = streaming optimise")

add_styled_heading("Comparatif Hebergeurs", 2)
add_table(
    ["Service", "Prix", "Stockage", "Bande passante", "Pour 500 videos*"],
    [
        ["Bunny.net", "~1-10 EUR/mois", "0.01 EUR/GB", "0.005 EUR/GB", "~5-15 EUR/mois"],
        ["Vimeo Plus", "84 EUR/an", "5GB/sem", "Fair use", "~180 EUR/an"],
        ["Vimeo Pro", "240 EUR/an", "20GB/sem", "Fair use", "~240 EUR/an"],
        ["Wistia Pro", "99 EUR/mois", "50 videos", "0.25 EUR/GB exces", "~1 200 EUR/an"],
    ],
    highlight_col=0,
)
add_body("*Estimation basee sur ~100GB de videos totales, trafic modere", italic=True, color=SABLE)

add_styled_heading("Bunny.net : La Solution Recommandee", 2)

add_table(
    ["Avantage", "Detail"],
    [
        ["Prix imbattable", "Minimum 1 EUR/mois, puis a l'usage"],
        ["Stockage", "0.01 EUR/GB/mois (~1 EUR pour 100GB)"],
        ["CDN", "0.005 EUR/GB (~5 EUR pour 1TB de streaming)"],
        ["Transcodage", "GRATUIT"],
        ["Transcription IA", "0.10 EUR/minute/langue"],
        ["119 PoPs mondiaux", "Performance globale"],
        ["Accessibilite", "Fonctionne dans les pays ou Vimeo est bloque"],
    ],
)

add_styled_heading("Pour Aurelia", 3)
add_bullet("500 videos (~100GB) = ~1 EUR/mois de stockage")
add_bullet("Streaming mensuel (estime 500GB) = ~2.50 EUR/mois")
add_body("Total estime : 5-15 EUR/mois vs 1 000 EUR/8 mois sur Uscreen", bold=True, color=TERRACOTTA)

add_styled_heading("Integration", 3)
add_bullet("Embed code compatible avec Kajabi, Circle, Skool")
add_bullet("Player personnalisable")
add_bullet("Protection par token (empeche le partage)")
add_bullet("Statistiques de visionnage")

add_styled_heading("Comment proceder ?", 2)
add_bullet("Extraire les videos de Uscreen (export ou telechargement manuel)")
add_bullet("Upload sur Bunny Stream")
add_bullet("Generer les liens d'embed")
add_bullet("Integrer dans la nouvelle plateforme")
add_body("Temps estime : 1-2 jours avec l'aide de Laurie", italic=True, color=CUIVRE)

# =============================================================================
# 5. ANALYSE COMPARATIVE DETAILLEE
# =============================================================================

doc.add_page_break()
add_styled_heading("5. Analyse Comparative Detaillee", 1)

# A. UX
add_styled_heading("A. Experience Utilisateur (UX)", 2)
add_table(
    ["Aspect", "Kajabi", "Circle", "Skool"],
    [
        ["Premiere impression", "Complexe", "Moderne", "Familier"],
        ["Courbe d'apprentissage", "Longue", "Moyenne", "Courte"],
        ["Navigation mobile", "Fragmentee", "Fluide", "Excellente"],
        ["Coherence interface", "Apps separees", "Unifiee", "Tres unifiee"],
        ["Personnalisation", "Templates", "CSS custom", "Limitee"],
    ],
)
add_winner("Skool (simplicite) / Circle (sophistication)")

# B. Cours
add_styled_heading("B. Fonctionnalites Cours", 2)
add_table(
    ["Fonctionnalite", "Kajabi", "Circle", "Skool"],
    [
        ["Modules structures", "Oui", "Oui", "Oui"],
        ["Drip content", "Oui", "Oui", "Oui (par niveau)"],
        ["Quiz/Evaluations", "Oui", "Non", "Non"],
        ["Certificats", "Oui", "Non", "Non"],
        ["Progression trackee", "Oui", "Oui", "Oui"],
        ["Multiples instructeurs", "Oui", "Oui", "Limite"],
        ["Cohorts/Cohortes", "Oui", "Oui", "Non"],
        ["Coaching 1:1", "Oui", "Oui", "Via calendrier"],
    ],
)
add_winner("Kajabi (nettement)")

# C. Communaute
add_styled_heading("C. Fonctionnalites Communaute", 2)
add_table(
    ["Fonctionnalite", "Kajabi", "Circle", "Skool"],
    [
        ["Feed style social", "Oui", "Oui", "Oui"],
        ["Sous-groupes", "Oui", "Oui (Spaces)", "Non"],
        ["Messagerie directe", "Oui", "Oui", "Oui"],
        ["Lives integres", "Oui", "Oui", "Via externe"],
        ["Gamification", "Basique", "Avancee", "Excellente"],
        ["Leaderboards", "Oui", "Oui", "Oui"],
        ["Profils membres", "Basique", "Riches", "Moyen"],
        ["Calendrier evenements", "Oui", "Oui", "Oui"],
    ],
)
add_winner("Circle / Skool (ex-aequo)")

# D. Marketing
add_styled_heading("D. Marketing et Vente", 2)
add_table(
    ["Fonctionnalite", "Kajabi", "Circle", "Skool"],
    [
        ["Landing pages", "Oui", "Non", "Non"],
        ["Funnels de vente", "Oui", "Non", "Non"],
        ["Email marketing", "Avance", "Basique", "Non"],
        ["Automations", "Avance", "Workflows", "Non"],
        ["Programme affiliation", "Oui", "Non", "Oui (40%)"],
        ["Checkout natif", "Oui", "Oui", "Oui"],
        ["Upsells", "Oui", "Non", "Non"],
    ],
)
add_winner("Kajabi (sans conteste)")

# E. Scalabilite
add_styled_heading("E. Scalabilite et Couts a Long Terme", 2)

add_body("Projection a 500 membres actifs :", bold=True)
add_table(
    ["Plateforme", "Cout fixe/an", "Frais variables", "Total estime"],
    [
        ["Uscreen actuel", "6 000 EUR", "500 x 1 EUR x 12 = 6 000 EUR", "12 000 EUR"],
        ["Kajabi Growth", "2 388 EUR", "0 EUR", "2 388 EUR"],
        ["Circle Business", "2 388 EUR", "~1% sur CA", "~3 000 EUR"],
        ["Skool Pro", "1 188 EUR", "2.9% sur CA", "~2 500 EUR"],
    ],
)

add_body("Projection a 2 000 membres actifs :", bold=True)
add_table(
    ["Plateforme", "Cout fixe/an", "Frais variables", "Total estime"],
    [
        ["Uscreen actuel", "6 000 EUR", "2000 x 1 EUR x 12 = 24 000 EUR", "30 000 EUR"],
        ["Kajabi Growth", "2 388 EUR", "0 EUR", "2 388 EUR"],
        ["Circle Business", "2 388 EUR", "~1% sur CA", "~4 000 EUR"],
        ["Skool Pro", "1 188 EUR", "2.9% sur CA", "~4 500 EUR"],
    ],
)

add_body("Conclusion : Toutes les alternatives sont MASSIVEMENT plus economiques qu'Uscreen.", bold=True, color=TERRACOTTA)

# =============================================================================
# 6. RECOMMANDATION FINALE
# =============================================================================

doc.add_page_break()
add_styled_heading("6. Recommandation Finale", 1)

# Reco principale
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Recommandation Principale : CIRCLE.SO (Plan Business)")
run.font.name = FONT_TITRE
run.font.size = Pt(15)
run.font.color.rgb = TERRACOTTA
run.bold = True

add_styled_heading("Pourquoi Circle ?", 2)
add_bullet("Migration assistee : Service gratuit pour les 500+ videos = gain de temps enorme", bold_prefix="1. ")
add_bullet("UX premium : Alignee avec son positionnement a 299 EUR", bold_prefix="2. ")
add_bullet("Communaute sophistiquee : Remplace avantageusement WhatsApp", bold_prefix="3. ")
add_bullet("Scalable sans frais par utilisateur : Plus jamais de limite de croissance", bold_prefix="4. ")
add_bullet("Multi-intervenants : Integration facile des 3 personnes formees", bold_prefix="5. ")
add_bullet("Delegation : Workflows pour automatiser et liberer Aurelia", bold_prefix="6. ")

add_styled_heading("Configuration recommandee", 2)
add_table(
    ["Element", "Solution", "Cout/an"],
    [
        ["Plateforme", "Circle Business (annuel)", "2 388 EUR"],
        ["Videos", "Bunny.net", "~120 EUR"],
        ["Email", "Mailerlite Pro", "~300 EUR"],
        ["Landing pages", "Site existant ou Carrd", "0-50 EUR"],
        ["TOTAL", "", "~2 800 EUR"],
    ],
)

add_body("Economie vs Uscreen actuel : ~4 200 EUR/an minimum (et plus si croissance)", bold=True, color=TERRACOTTA)

# Alternatives
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
run = p.add_run("─" * 60)
run.font.color.rgb = SABLE
run.font.size = Pt(8)

add_styled_heading("Alternative 1 : KAJABI (Si Marketing Prioritaire)", 2)
add_bullet("Aurelia veut TOUT dans une seule plateforme")
add_bullet("Les funnels de vente sont essentiels")
add_bullet("L'email marketing avance est crucial")
add_bullet("Le multilingue FR/EN est prioritaire a court terme")
add_body("Attention : UX moins fluide, communaute limited, plus cher", italic=True, color=CUIVRE)

add_styled_heading("Alternative 2 : SKOOL (Si Communaute Only)", 2)
add_bullet("La communaute est LE coeur du business")
add_bullet("Le budget est ultra-serre")
add_bullet("La simplicite absolue est souhaitee")
add_body("Attention : Pas adapte a un positionnement premium, migration 100% manuelle", italic=True, color=CUIVRE)

# =============================================================================
# 7. PLAN D'ACTION MIGRATION
# =============================================================================

doc.add_page_break()
add_styled_heading("7. Plan d'Action Migration", 1)

add_styled_heading("Phase 1 : Preparation (Semaine 1-2)", 2)
add_bullet("Extraire la liste complete des videos Uscreen")
add_bullet("Telecharger toutes les videos (ou demander export)")
add_bullet("Creer compte Bunny.net et uploader les videos")
add_bullet("Contacter Circle pour activer la migration gratuite")
add_bullet("Documenter la structure actuelle des programmes")

add_styled_heading("Phase 2 : Configuration (Semaine 3-4)", 2)
add_bullet("Setup Circle Business avec le branding d'Aurelia")
add_bullet("Creer les Spaces (communaute, chaque programme)")
add_bullet("Configurer les niveaux d'acces")
add_bullet("Integrer les videos Bunny.net")
add_bullet("Configurer les paiements (Stripe)")
add_bullet("Former Laurie a l'administration")

add_styled_heading("Phase 3 : Migration Membres (Semaine 5-6)", 2)
add_bullet("Utiliser le service de migration paiements de Circle")
add_bullet("Communiquer aux membres le changement")
add_bullet("Periode de double acces (ancien + nouveau)")
add_bullet("Support actif pendant la transition")
add_bullet("Fermer Uscreen une fois migration complete")

add_styled_heading("Phase 4 : Optimisation (Semaine 7-8)", 2)
add_bullet("Configurer les automations/workflows")
add_bullet("Mettre en place la gamification")
add_bullet("Integrer les 3 intervenants")
add_bullet("Former Aurelia au tableau de bord")
add_bullet("KPIs de suivi engagement")

# =============================================================================
# PIED DE PAGE
# =============================================================================

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─────────────────")
run.font.color.rgb = SABLE
run.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Catherine Selosse — Consulting Strategique")
run.font.name = FONT_TITRE
run.font.size = Pt(10)
run.font.color.rgb = TERRE
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Fevrier 2026")
run.font.name = FONT_CORPS
run.font.size = Pt(9)
run.font.color.rgb = SABLE

# === SAUVEGARDE ===
output_path = os.path.expanduser("~/Downloads/Comparatif_Plateformes_Aurelia_DELSOL.docx")
doc.save(output_path)
print(f"Document genere : {output_path}")
